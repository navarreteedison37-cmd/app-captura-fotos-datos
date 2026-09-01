from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os

class WordGenerator:
    """Genera documentos Word con informes"""
    
    def __init__(self):
        self.output_dir = 'reportes'
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def generar(self, registro):
        """Genera un informe Word basado en el tipo de registro"""
        try:
            doc = Document()
            
            # Título
            titulo = doc.add_heading('INFORME OFICIAL', 0)
            titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Tipo de informe
            tipo = "INSPECCIÓN LABORAL" if registro['tipo_informe'] == 'inspeccion_laboral' else "CONTROL DE INGRESOS"
            subtitulo = doc.add_heading(tipo, level=1)
            subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Fecha de generación
            fecha_gen = doc.add_paragraph()
            fecha_gen.add_run(f"Fecha de Generación: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
            fecha_gen.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            doc.add_paragraph()  # Espacio
            
            # Datos del Comprobante
            doc.add_heading('DATOS DEL COMPROBANTE', level=2)
            comprobante = registro['comprobante']
            self._agregar_tabla_datos(doc, [
                ("Número de Carta:", comprobante.get('numero_carta', 'N/A')),
                ("Recepción:", comprobante.get('recepcion', 'N/A')),
                ("Vínculo:", comprobante.get('vinculo', 'N/A')),
                ("Fecha:", comprobante.get('fecha', 'N/A')),
                ("Hora:", comprobante.get('hora', 'N/A')),
            ])
            
            doc.add_paragraph()  # Espacio
            
            # Datos del Documento
            doc.add_heading('DATOS DEL DOCUMENTO', level=2)
            documento = registro['documento']
            self._agregar_tabla_datos(doc, [
                ("Número de Acta:", documento.get('numero_acta', 'N/A')),
                ("RUC:", documento.get('ruc', 'N/A')),
                ("Nombre/Razón Social:", documento.get('nombre_razon_social', 'N/A')),
            ])
            
            doc.add_paragraph()  # Espacio
            
            # Observaciones
            doc.add_heading('OBSERVACIONES', level=2)
            obs = doc.add_paragraph()
            obs.add_run("[Agregar observaciones según corresponda]")
            
            doc.add_paragraph()  # Espacio
            
            # Pie
            doc.add_paragraph()
            pie = doc.add_paragraph()
            pie.add_run(f"Informe generado automáticamente - {datetime.now().strftime('%d/%m/%Y')}").font.size = Pt(9)
            pie.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Guardar archivo
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            archivo = os.path.join(self.output_dir, f'informe_{tipo.replace(" ", "_")}_{timestamp}.docx')
            doc.save(archivo)
            
            return archivo
        except Exception as e:
            raise Exception(f"Error generando Word: {str(e)}")
    
    def _agregar_tabla_datos(self, doc, datos):
        """Agrega una tabla con datos de label-value"""
        table = doc.add_table(rows=len(datos) + 1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Headers
        table.rows[0].cells[0].text = "Campo"
        table.rows[0].cells[1].text = "Valor"
        
        # Datos
        for i, (label, valor) in enumerate(datos, 1):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(valor)
